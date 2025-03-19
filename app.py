import streamlit as st
import PyPDF2
import os
import io
import docx
import unicodedata
from fuzzywuzzy import fuzz
from fpdf import FPDF
from docx import Document
from crewai_tools import SerperDevTool
from crewai import Agent, Task, Crew, Process
from dotenv import load_dotenv
from openai import OpenAI

# ✅ Load API Keys
load_dotenv(override=True)

# ✅ Retrieve API Keys
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
SERPER_API_KEY = os.getenv("SERPER_API_KEY")

# ✅ Validate API Keys
if not OPENAI_API_KEY or not SERPER_API_KEY:
    st.error("❌ API keys missing! Please check your .env file.")
    raise ValueError("Missing API keys.")

# ✅ Debugging in console only
print(f"✅ OPENAI_API_KEY loaded: {OPENAI_API_KEY[:5]}*****")
print(f"✅ SERPER_API_KEY loaded: {SERPER_API_KEY[:5]}*****")

# ✅ Set API keys as environment variables
os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
os.environ["SERPER_API_KEY"] = SERPER_API_KEY


### ✅ Extract Text from Different File Types ###
def extract_text_from_pdf(uploaded_file):
    """Extract text from a PDF file uploaded via Streamlit."""
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        return text.strip() if text else "❌ No text found in PDF."
    except Exception as e:
        return f"❌ Error extracting text from PDF: {e}"


def extract_text_from_docx(uploaded_file):
    """Extract text from a Word document (DOCX)."""
    try:
        doc = Document(uploaded_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip() if text else "❌ No text found in DOCX."
    except Exception as e:
        return f"❌ Error extracting text from DOCX: {e}"


def extract_text_from_txt(uploaded_file):
    """Extract text from a plain text file (TXT)."""
    try:
        return uploaded_file.read().decode("utf-8").strip()
    except Exception as e:
        return f"❌ Error extracting text from TXT: {e}"


def extract_text(uploaded_file):
    """Extract text based on file type from an UploadedFile object."""
    if not uploaded_file:
        return None

    file_extension = uploaded_file.name.split(".")[-1].lower()

    if file_extension == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == "docx":
        return extract_text_from_docx(uploaded_file)
    elif file_extension == "txt":
        return extract_text_from_txt(uploaded_file)
    else:
        return "❌ Unsupported file format. Please upload a PDF, DOCX, or TXT file."


### ✅ Extract Keywords Using OpenAI GPT ###
def extract_keywords_with_gpt(text):
    """Use OpenAI GPT to extract must-have and good-to-have skills from job description."""
    client = OpenAI(api_key=OPENAI_API_KEY)

    try:
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "Extract must-have and good-to-have skills from this job description."},
                {"role": "user", "content": text}
            ],
            max_tokens=150,
            temperature=0.2
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return "⚠️ Unable to extract job keywords due to API error."


### ✅ Extract Keywords Using Serper ###
def extract_keywords_with_serper(job_description):
    """Use Serper API to search for job-related keywords."""
    tool = SerperDevTool()

    try:
        result = tool.run()
        return result if isinstance(result, str) else str(result)
    except Exception:
        return None  # ✅ Handle failure gracefully


### ✅ Match Scoring System ###
def calculate_match_score(resume_text, job_desc_text):
    """Calculate similarity match score using fuzzy keyword comparison."""
    return fuzz.token_set_ratio(resume_text, job_desc_text)


### ✅ Save as DOCX ###
def save_as_docx(text):
    """Save the optimized resume as a DOCX file."""
    doc = Document()
    doc.add_paragraph(text)

    docx_filename = "optimized_resume.docx"
    doc.save(docx_filename)
    
    return docx_filename


### ✅ Save as PDF ###
def save_as_pdf(text, filename="optimized_resume.pdf"):
    """Save the optimized resume as a PDF file, ensuring it is free from unsupported Unicode characters."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # ✅ Normalize text to remove unsupported Unicode characters
    safe_text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")

    pdf.multi_cell(190, 10, safe_text)
    pdf.output(filename, "F")

    return filename

### ✅ Streamlit UI ###
st.title("🔍 AI Resume Matcher & Optimizer")
st.write("Upload your resume and job description to get an optimized resume with keyword matching and a match score.")

# Upload Resume
uploaded_resume = st.file_uploader("📄 Upload your **RESUME** (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
resume_text = extract_text(uploaded_resume) if uploaded_resume else None

# Upload Job Description
uploaded_job_desc = st.file_uploader("📑 Upload the **JOB DESCRIPTION** (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
job_description_text = extract_text(uploaded_job_desc) if uploaded_job_desc else None

# Display Extracted Text
if resume_text:
    st.text_area("📄 **Extracted Resume Text:**", resume_text, height=200)

if job_description_text:
    st.text_area("📑 **Extracted Job Description Text:**", job_description_text, height=200)


if st.button("✨ Optimize Resume"):
    if not resume_text or not job_description_text:
        st.error("❌ Please upload both Resume and Job Description files!")
    else:
        # Extract job keywords with GPT
        job_keywords_gpt = extract_keywords_with_gpt(job_description_text)
        
        # Extract job keywords with Serper, but **skip error messages**
        job_keywords_serper = extract_keywords_with_serper(job_description_text)
        
        # ✅ Only include valid Serper output
        combined_keywords = job_keywords_gpt
        if job_keywords_serper:
            combined_keywords += f"\n\n{job_keywords_serper}"

        st.write("🔑 **Extracted Job Keywords (GPT & Serper):**")
        st.text_area("", combined_keywords, height=100)

        # Calculate match score
        match_score = calculate_match_score(resume_text, job_description_text)
        st.metric("🔍 Resume Match Score", f"{match_score}%")

        # Optimize Resume (without raw errors)
        optimized_resume = f"### Optimized Resume\n\n{resume_text}"

        st.subheader("✅ **Optimized Resume:**")
        st.text_area("📌 Updated Resume:", optimized_resume, height=250)

        docx_filename = save_as_docx(optimized_resume)
        pdf_filename = save_as_pdf(optimized_resume)

        st.download_button("📥 Download DOCX", open(docx_filename, "rb"), file_name="Optimized_Resume.docx")
        st.download_button("📥 Download PDF", open(pdf_filename, "rb"), file_name="Optimized_Resume.pdf")

st.write("🚀 Built with OpenAI, Serper, Streamlit & NLP for smarter job applications!")
